# import fitz  # PyMuPDF
# import docx
# import requests
# from langchain.text_splitter import RecursiveCharacterTextSplitter

# def download_and_read_file(url: str) -> str:
#     response = requests.get(url)
#     if url.endswith(".pdf"):
#         with open("temp.pdf", "wb") as f:
#             f.write(response.content)
#         doc = fitz.open("temp.pdf")
#         return "\n".join([page.get_text() for page in doc])
#     elif url.endswith(".docx"):
#         with open("temp.docx", "wb") as f:
#             f.write(response.content)
#         doc = docx.Document("temp.docx")
#         return "\n".join([para.text for para in doc.paragraphs])
#     else:
#         raise ValueError("Unsupported file type.")

# def chunk_text(text: str) -> list:
#     splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
#     return splitter.split_text(text)
import fitz  # PyMuPDF
import docx
import requests
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
from langchain.text_splitter import RecursiveCharacterTextSplitter
import numpy as np

semantic_embedder  = SentenceTransformer("all-MiniLM-L6-v2")

def download_and_read_file(url: str) -> str:
    response = requests.get(url)
    if url.endswith(".pdf"):
        with open("temp.pdf", "wb") as f:
            f.write(response.content)
        doc = fitz.open("temp.pdf")
        return "\n".join([page.get_text() for page in doc])
    elif url.endswith(".docx"):
        with open("temp.docx", "wb") as f:
            f.write(response.content)
        doc = docx.Document("temp.docx")
        return "\n".join([para.text for para in doc.paragraphs])
    else:
        raise ValueError("Unsupported file type.")

def chunk_text(text: str) -> list:
    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    return splitter.split_text(text)
# def chunk_text(text: str, similarity_threshold=0.5, max_chunk_len=1000):
#     sentences = [s.strip() for s in text.split('.') if s.strip()]
#     embeddings = semantic_embedder .encode(sentences)
    
#     chunks, current_chunk, current_vecs = [], [], []
#     for i, (sent, vec) in enumerate(zip(sentences, embeddings)):
#         if not current_chunk:
#             current_chunk.append(sent)
#             current_vecs.append(vec)
#             continue

#         similarity = cosine_similarity([vec], [np.mean(current_vecs, axis=0)])[0][0]

#         if similarity >= similarity_threshold and len(" ".join(current_chunk)) < max_chunk_len:
#             current_chunk.append(sent)
#             current_vecs.append(vec)
#         else:
#             chunks.append(". ".join(current_chunk))
#             current_chunk = [sent]
#             current_vecs = [vec]

#     if current_chunk:
#         chunks.append(". ".join(current_chunk))

#     return chunks
